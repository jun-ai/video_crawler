"""视频词汇解读 → 学习卡片导出 (PDF / Word)

参考版式: 谷爱凌学习卡片 PDF (绿色表头 + 行交替 + 单词列粗体斜体音标 + 例句 EN/CN 上下)
在此基础上的提升:
- 加封面页 (视频标题 + 日期 + 统计)
- 板块标题 banner (居中加粗大字)
- 难度/词频 badge
- 例句区左侧竖线
- 品牌主色 #4DA06C

支持两种字段范围:
- compact: 对标参考 PDF (一页约 7 条)
- full: 加难度/词频/英英释义/其他词性
"""

from __future__ import annotations

import io
import os
import re
import sys
import datetime
from typing import Iterable, List, Optional, Tuple

# reportlab (PDF)
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, CondPageBreak, KeepTogether, Flowable,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.pdfmetrics import registerFontFamily

# python-docx (Word)
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ==================== 设计 token ====================

BRAND_GREEN = colors.HexColor('#4DA06C')           # 主品牌色
BRAND_GREEN_LIGHT = colors.HexColor('#F0F9F3')      # 行交替浅绿
BRAND_GREEN_SUBTLE = colors.HexColor('#E6F4EC')     # banner 背景
ACCENT_PURPLE = colors.HexColor('#F5F0FF')          # 表达卡紫
ACCENT_PINK = colors.HexColor('#FFF0F5')            # 表达卡粉
TEXT_PRIMARY = colors.HexColor('#1E293B')
TEXT_SECONDARY = colors.HexColor('#64748B')
TEXT_MUTED = colors.HexColor('#94A3B8')
BORDER_COLOR = colors.HexColor('#E2E8F0')

DIFFICULTY_COLORS = [
    colors.HexColor('#94A3B8'),  # 1 入门 灰
    colors.HexColor('#60A5FA'),  # 2 简单 蓝
    colors.HexColor('#4DA06C'),  # 3 中等 绿 (品牌)
    colors.HexColor('#F59E0B'),  # 4 进阶 橙
    colors.HexColor('#EF4444'),  # 5 困难 红
]

# Hex string 版本, DOCX 走 OXML 写 hex, 不需要 colors.HexColor 对象
# PDF 仍用上面的 HexColor, DOCX 用下面这一份 — 同一组色值, 单一来源
# 注意: 不带 # (跟现有 _set_cell_shading / _add_para_run 约定一致,
#       那些函数把 hex 直接喂给 RGBColor.from_string 和 w:fill 槽位)
BRAND_GREEN_HEX = '4DA06C'
BRAND_GREEN_LIGHT_HEX = 'F0F9F3'
BRAND_GREEN_SUBTLE_HEX = 'E6F4EC'
ACCENT_PURPLE_HEX = 'F5F0FF'
ACCENT_PINK_HEX = 'FFF0F5'
TEXT_PRIMARY_HEX = '1E293B'
TEXT_SECONDARY_HEX = '64748B'
TEXT_MUTED_HEX = '94A3B8'
BORDER_COLOR_HEX = 'E2E8F0'
DIFFICULTY_COLORS_HEX = [
    '94A3B8',  # 1 入门 灰
    '60A5FA',  # 2 简单 蓝
    '4DA06C',  # 3 中等 绿 (品牌)
    'F59E0B',  # 4 进阶 橙
    'EF4444',  # 5 困难 红
]

OUTPUT_PAGE_SIZE = A4
PAGE_MARGIN_TOP = 28 * mm      # 留给页眉
PAGE_MARGIN_BOTTOM = 22 * mm   # 留给页脚
PAGE_MARGIN_LEFT = 16 * mm
PAGE_MARGIN_RIGHT = 16 * mm


# ==================== 字体解析 (跨平台 fallback) ====================

_FONT_REGISTERED = False
_FONT_REGULAR = 'NotoSansSC'
_FONT_BOLD = 'NotoSansSC-Bold'
# 音标字体: Microsoft YaHei / NotoSansSC 都缺 IPA 字符 (ˈ ɜː ɪ ʃ ŋ ...),
# 导致 /ˈfɜːrnɪʃɪŋ/ 被渲染成 // f rn ŋ// (丢字).
# 专门注册一个 IPA-capable 字体, 音标段落单独用它
_FONT_IPA = 'IPAFONT'


def _resolve_font_files() -> Tuple[str, Optional[str]]:
    """按优先级查找中文字体: 项目打包 → Windows 系统 → Linux 系统

    返回 (regular_path, bold_path)
    bold_path 可能为 None (用 regular 顶替)
    """
    # 1. 项目打包字体 (跨平台一致, 生产推荐)
    pkg_dir = os.path.join(os.path.dirname(__file__), '..', 'static', 'fonts')
    candidates_pkg = [
        ('NotoSansSC-Regular.otf', 'NotoSansSC-Bold.otf'),
        ('NotoSansSC-Regular.ttf', 'NotoSansSC-Bold.ttf'),
        ('SourceHanSansSC-Regular.otf', 'SourceHanSansSC-Bold.otf'),
    ]
    for reg_name, bold_name in candidates_pkg:
        reg_path = os.path.join(pkg_dir, reg_name)
        bold_path = os.path.join(pkg_dir, bold_name)
        if os.path.exists(reg_path):
            return reg_path, bold_path if os.path.exists(bold_path) else None

    # 2. Windows 系统
    if sys.platform == 'win32':
        win_dir = r'C:\Windows\Fonts'
        # 微软雅黑有专门的 bold
        msyh_reg = os.path.join(win_dir, 'msyh.ttc')
        msyh_bold = os.path.join(win_dir, 'msyhbd.ttc')
        if os.path.exists(msyh_reg):
            return msyh_reg, msyh_bold if os.path.exists(msyh_bold) else None
        # 黑体 (无 bold variant)
        simhei = os.path.join(win_dir, 'simhei.ttf')
        if os.path.exists(simhei):
            return simhei, None

    # 3. Linux 系统 (阿里云)
    else:
        linux_candidates = [
            '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/wqy-microhei/wqy-microhei.ttc',
            '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
        ]
        for path in linux_candidates:
            if os.path.exists(path):
                return path, None

    raise RuntimeError(
        '未找到中文字体。请把 NotoSansSC-Regular.otf + NotoSansSC-Bold.otf '
        '放到 backend/app/static/fonts/'
    )


def _resolve_ipa_font_file() -> Optional[str]:
    """查找 IPA 音标字体: 项目打包 → Windows → Linux

    Microsoft YaHei / Noto Sans SC 都不含 IPA 字符 (ˈ ɜ ː ɪ ʃ ʊ ə ...),
    音标段落需要单独用 IPA-capable 字体渲染.
    """
    # 1. 项目打包
    pkg_dir = os.path.join(os.path.dirname(__file__), '..', 'static', 'fonts')
    for name in ('DejaVuSans.ttf', 'NotoSans-Regular.ttf', 'SourceSansPro-Regular.ttf'):
        p = os.path.join(pkg_dir, name)
        if os.path.exists(p):
            return p

    # 2. Windows: Segoe UI / Calibri / Arial 都含 IPA
    if sys.platform == 'win32':
        win_dir = r'C:\Windows\Fonts'
        for name in ('segoeui.ttf', 'calibri.ttf', 'arial.ttf'):
            p = os.path.join(win_dir, name)
            if os.path.exists(p):
                return p

    # 3. Linux: DejaVu Sans (几乎所有发行版自带)
    else:
        for p in [
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            '/usr/share/fonts/dejavu/DejaVuSans.ttf',
            '/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf',
        ]:
            if os.path.exists(p):
                return p
    return None


def _ensure_font_registered():
    """注册中文字体到 reportlab (只注册一次)"""
    global _FONT_REGISTERED
    if _FONT_REGISTERED:
        return

    reg_path, bold_path = _resolve_font_files()

    # .ttc 是字体集合, 需要指定 subfontIndex
    is_ttc = reg_path.lower().endswith('.ttc')
    if is_ttc:
        pdfmetrics.registerFont(TTFont(_FONT_REGULAR, reg_path, subfontIndex=0))
        if bold_path:
            pdfmetrics.registerFont(TTFont(_FONT_BOLD, bold_path, subfontIndex=0))
        else:
            # .ttc 里通常第 0 个就是 regular, bold 找不到就退化为 regular
            try:
                pdfmetrics.registerFont(TTFont(_FONT_BOLD, reg_path, subfontIndex=1))
            except Exception:
                pass
    else:
        pdfmetrics.registerFont(TTFont(_FONT_REGULAR, reg_path))
        if bold_path:
            pdfmetrics.registerFont(TTFont(_FONT_BOLD, bold_path))

    # 注册 family 让 <b> 自动切到 bold
    registerFontFamily(
        _FONT_REGULAR,
        normal=_FONT_REGULAR,
        bold=_FONT_BOLD if bold_path or is_ttc else _FONT_REGULAR,
        italic=_FONT_REGULAR,        # 中文字体没有真正的 italic, 用 regular 顶替
        boldItalic=_FONT_BOLD if bold_path or is_ttc else _FONT_REGULAR,
    )

    # 注册 IPA 字体 (可选, 没找到就降级到主字体, 音标会显示不全)
    ipa_path = _resolve_ipa_font_file()
    if ipa_path:
        try:
            pdfmetrics.registerFont(TTFont(_FONT_IPA, ipa_path))
        except Exception:
            pass

    _FONT_REGISTERED = True


# ==================== 数据分类与字段 ====================

def _categorize(interpretations: Iterable) -> dict:
    """把 interpretation 列表按 category 分组

    返回:
        {
            'words': [...],         # category == 'word'
            'phrases': [...],       # category == 'phrase'
            'expressions': [...],   # category in ('grammar', 'idiom')
        }
    按 sequence 排序
    """
    groups = {'words': [], 'phrases': [], 'expressions': []}
    for item in interpretations:
        cat = (getattr(item, 'category', '') or '').lower()
        if cat == 'word':
            groups['words'].append(item)
        elif cat == 'phrase':
            groups['phrases'].append(item)
        elif cat in ('grammar', 'idiom'):
            groups['expressions'].append(item)
    for k in groups:
        groups[k].sort(key=lambda x: getattr(x, 'sequence', 0) or 0)
    return groups


def _compute_vocab_level(items: Iterable) -> str:
    """根据 frequency_rank 平均值估算词汇等级

    规则:
        avg <= 1000   → 'CET-4'
        avg <= 3000   → 'CET-6'
        avg <= 5000   → 'IELTS 6.5-7.0'
        avg <= 8000   → 'IELTS 7.5+'
        avg > 8000    → 'Advanced / 高级'
        无 frequency_rank → '—'
    """
    ranks = [getattr(it, 'frequency_rank', None) for it in items
             if getattr(it, 'frequency_rank', None) is not None]
    if not ranks:
        return '—'
    avg = sum(ranks) / len(ranks)
    if avg <= 1000:
        return 'CET-4'
    if avg <= 3000:
        return 'CET-6'
    if avg <= 5000:
        return 'IELTS 6.5-7.0'
    if avg <= 8000:
        return 'IELTS 7.5+'
    return 'Advanced / 高级'


def _compute_difficulty_distribution(items: Iterable) -> dict:
    """统计 difficulty 1-5 的分布

    返回 {1: n1, 2: n2, 3: n3, 4: n4, 5: n5}, 缺的 key 默认 0
    """
    dist = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
    for it in items:
        d = getattr(it, 'difficulty', None)
        if d is None:
            continue
        try:
            d_int = int(d)
        except (TypeError, ValueError):
            continue
        if 1 <= d_int <= 5:
            dist[d_int] += 1
    return dist


def _compute_reading_minutes(words_n: int, phrases_n: int, expressions_n: int) -> int:
    """估算阅读/学习时长 (分钟, 向上取整, 至少 1 分钟)

    启发式:
        - 单词: 30 秒 (读音标 + 释义 + 例句)
        - 短语: 60 秒 (比单词长, 需要理解搭配)
        - 表达: 90 秒 (结构解析 + 举一反三 + 场景)
    """
    total_seconds = words_n * 30 + phrases_n * 60 + expressions_n * 90
    return max(1, (total_seconds + 59) // 60)


def _safe_text(s, default='') -> str:
    if s is None:
        return default
    return str(s).strip()


def _escape_xml(s) -> str:
    """转义 Paragraph 里的特殊字符"""
    if s is None:
        return ''
    s = str(s)
    return s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


# emoji / 符号清理 — Microsoft YaHei / NotoSansSC 都缺 emoji glyph,
# ✨(U+2728) 之类的字符在 PDF 里会渲染成 NULL / 方框 □
_EMOJI_RE = re.compile(
    '['
    '\U0001F300-\U0001FAFF'   # symbols & pictographs + extensions
    '\U00002600-\U000027BF'   # misc symbols & dingbats (✨ = U+2728)
    '\U0001F900-\U0001F9FF'   # supplemental symbols
    '\u2000-\u200F'           # zero-width / dash/punctuation specials
    '\u2028-\u202F'           # line/para sep, hyphens
    ']',
    flags=re.UNICODE,
)


def _sanitize_title(s) -> str:
    """剔除字体不支持的 emoji / 特殊符号, 压缩多余空白

    用于视频标题在 PDF / DOCX 显示前清理, 避免 ✨ 等字符渲染成方框.
    """
    if not s:
        return ''
    cleaned = _EMOJI_RE.sub(' ', str(s))
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned


# ==================== PDF Paragraph styles ====================

def _make_styles():
    """构造常用 ParagraphStyle"""
    _ensure_font_registered()
    base = dict(fontName=_FONT_REGULAR, textColor=TEXT_PRIMARY, leading=16)
    # IPA 字体是否注册了 (没注册就降级到主字体, 音标会丢字但不会崩)
    ipa_available = _FONT_IPA in pdfmetrics.getRegisteredFontNames()
    ipa_font_name = _FONT_IPA if ipa_available else _FONT_REGULAR
    return {
        'h1_cover': ParagraphStyle('h1_cover', **{**base,
            'fontName': _FONT_BOLD, 'fontSize': 28, 'leading': 36,
            'alignment': TA_CENTER, 'textColor': BRAND_GREEN, 'spaceAfter': 8}),
        'h2_cover_sub': ParagraphStyle('h2_cover_sub', **{**base,
            'fontSize': 13, 'leading': 18, 'alignment': TA_CENTER,
            'textColor': TEXT_SECONDARY, 'spaceAfter': 4}),
        'cover_stat': ParagraphStyle('cover_stat', **{**base,
            'fontSize': 11, 'leading': 16, 'alignment': TA_CENTER,
            'textColor': TEXT_SECONDARY}),
        'section_banner': ParagraphStyle('section_banner', **{**base,
            'fontName': _FONT_BOLD, 'fontSize': 18, 'leading': 26,
            'alignment': TA_CENTER, 'textColor': BRAND_GREEN,
            'spaceBefore': 4, 'spaceAfter': 12}),
        'table_header': ParagraphStyle('table_header', **{**base,
            'fontName': _FONT_BOLD, 'fontSize': 11, 'leading': 14,
            'textColor': colors.white}),
        'word_text': ParagraphStyle('word_text', **{**base,
            'fontName': _FONT_BOLD, 'fontSize': 12, 'leading': 15,
            'textColor': BRAND_GREEN}),
        # 音标段落: 用 IPA 字体, 否则 ˈ ɜː ɪ ʃ 等字符会被丢
        'phonetic_ipa': ParagraphStyle('phonetic_ipa', **{**base,
            'fontName': ipa_font_name, 'fontSize': 10, 'leading': 13,
            'textColor': TEXT_SECONDARY}),
        # 小字 metadata (难度/词频/英英等): 含中文, 用主字体
        'phonetic': ParagraphStyle('phonetic', **{**base,
            'fontSize': 10, 'leading': 13, 'textColor': TEXT_SECONDARY}),
        'pos_cn': ParagraphStyle('pos_cn', **{**base,
            'fontSize': 11, 'leading': 15, 'textColor': TEXT_PRIMARY, 'spaceAfter': 2}),
        'example_en': ParagraphStyle('example_en', **{**base,
            'fontSize': 10, 'leading': 13, 'textColor': TEXT_PRIMARY,
            'leftIndent': 8, 'spaceAfter': 0}),
        'example_cn': ParagraphStyle('example_cn', **{**base,
            'fontSize': 10, 'leading': 13, 'textColor': TEXT_SECONDARY,
            'leftIndent': 8, 'spaceAfter': 2}),
        'label_small': ParagraphStyle('label_small', **{**base,
            'fontName': _FONT_BOLD, 'fontSize': 9, 'leading': 12,
            'textColor': TEXT_MUTED}),
        'meta_text': ParagraphStyle('meta_text', **{**base,
            'fontSize': 10, 'leading': 13, 'textColor': TEXT_SECONDARY}),
        'expression_title': ParagraphStyle('expression_title', **{**base,
            'fontName': _FONT_BOLD, 'fontSize': 12, 'leading': 16,
            'textColor': TEXT_PRIMARY, 'spaceAfter': 4}),
        'no': ParagraphStyle('no', **{**base,
            'fontSize': 10, 'leading': 13, 'alignment': TA_CENTER,
            'textColor': TEXT_MUTED}),
    }


# ==================== PDF 构建 ====================

def _on_page(canvas, doc, material_title):
    """页眉页脚 (每页都画)"""
    canvas.saveState()
    width, height = OUTPUT_PAGE_SIZE

    # 页眉: 左品牌 / 中板块名 (空, 每个板块自己画 banner) / 右视频名
    canvas.setFont(_FONT_REGULAR, 9)
    canvas.setFillColor(TEXT_MUTED)
    canvas.drawString(PAGE_MARGIN_LEFT, height - 14 * mm, 'Linyu 学习卡')
    # 视频名右对齐 (截断到 40 字符)
    title = material_title or ''
    if len(title) > 40:
        title = title[:39] + '…'
    canvas.drawRightString(width - PAGE_MARGIN_RIGHT, height - 14 * mm, title)
    # 页眉底部细线
    canvas.setStrokeColor(BORDER_COLOR)
    canvas.setLineWidth(0.4)
    canvas.line(PAGE_MARGIN_LEFT, height - 17 * mm,
                width - PAGE_MARGIN_RIGHT, height - 17 * mm)

    # 页脚: 左来源 / 右页码
    canvas.setFont(_FONT_REGULAR, 9)
    canvas.setFillColor(TEXT_MUTED)
    canvas.drawString(PAGE_MARGIN_LEFT, 10 * mm, 'From Linyu')
    page_num = canvas.getPageNumber()
    canvas.drawRightString(width - PAGE_MARGIN_RIGHT, 10 * mm, f'— {page_num:02d} —')

    canvas.restoreState()


def _build_cover(material_title, stats, today_str, styles, items=None):
    """封面页 flowables

    Phase C: 在装饰线 + 日期之后, 统计行之前, 插入 3 项数据可视化
    - 阅读时长估算 ("≈ 25 分钟")
    - 难度分布条 (5 色按比例)
    - 词汇等级 pill ("CET-6 / IELTS 7.0+")
    items: 全部 interpretations 列表 (用于计算 vocab level / difficulty distribution)
    """
    flow = []
    flow.append(Spacer(1, 50 * mm))
    flow.append(Paragraph('Linyu 学习卡', styles['h2_cover_sub']))
    flow.append(Spacer(1, 8 * mm))
    flow.append(Paragraph(_escape_xml(material_title or '视频学习卡片'), styles['h1_cover']))
    flow.append(Spacer(1, 4 * mm))
    # 装饰横线
    hr = Table([['']], colWidths=[60 * mm], rowHeights=[2])
    hr.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), BRAND_GREEN),
        ('LINEBELOW', (0, 0), (-1, -1), 0, colors.transparent),
    ]))
    # 居中: 用一个 3 列 table 模拟
    centered_hr = Table([[Spacer(1, 1), hr, Spacer(1, 1)]],
                       colWidths=[None, 60 * mm, None])
    centered_hr.setStyle(TableStyle([('ALIGN', (1, 0), (1, 0), 'CENTER')]))
    flow.append(centered_hr)
    flow.append(Spacer(1, 8 * mm))
    flow.append(Paragraph(today_str, styles['h2_cover_sub']))
    flow.append(Spacer(1, 16 * mm))

    # ===== Phase C: 3 项数据可视化 =====
    if items:
        # 1. 阅读时长
        mins = _compute_reading_minutes(
            stats['words'], stats['phrases'], stats['expressions'])
        flow.append(Paragraph(
            f'<font color="#94A3B8">预计学习时长</font>  '
            f'<font color="#1E293B"><b>≈ {mins} 分钟</b></font>',
            ParagraphStyle('cover_meta', parent=styles['cover_stat'],
                           fontSize=12, leading=18)
        ))
        flow.append(Spacer(1, 4 * mm))

        # 2. 难度分布条 (5 段彩色横条, 宽度按比例)
        diff_dist = _compute_difficulty_distribution(items)
        total = sum(diff_dist.values())
        if total > 0:
            bar_w = 100  # mm 总宽
            cell_widths = []
            cell_colors = []
            for d in range(1, 6):
                w = max(2, bar_w * diff_dist[d] / total)  # 最少 2mm, 保证可见
                cell_widths.append(w * mm)
                cell_colors.append(DIFFICULTY_COLORS[d - 1])
            bar_table = Table([[''] * 5], colWidths=cell_widths, rowHeights=[4 * mm])
            bg_styles = [('BACKGROUND', (i, 0), (i, 0), cell_colors[i])
                         for i in range(5)]
            sep_styles = [('LINEAFTER', (i, 0), (i, 0), 0.5, colors.white)
                          for i in range(4)]
            common_styles = [
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]
            bar_table.setStyle(TableStyle(bg_styles + common_styles + sep_styles))
            centered_bar = Table([[Spacer(1, 1), bar_table, Spacer(1, 1)]],
                                 colWidths=[None, bar_w * mm, None])
            centered_bar.setStyle(TableStyle([
                ('ALIGN', (1, 0), (1, 0), 'CENTER'),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ]))
            flow.append(centered_bar)
            flow.append(Spacer(1, 2 * mm))
            # 难度条下加图例 (难 1 → 难 5)
            # 注: reportlab <font color="..."> 需要带 #, DIFFICULTY_COLORS_HEX 是不带的常量
            legend = Paragraph(
                '<font color="#94A3B8" size="9">'
                '简单  ' +
                ''.join([f'<font color="#{DIFFICULTY_COLORS_HEX[i]}">●</font>'
                          for i in range(5)]) +
                '  困难</font>',
                ParagraphStyle('cover_legend', parent=styles['cover_stat'],
                               fontSize=9, alignment=TA_CENTER)
            )
            flow.append(legend)
            flow.append(Spacer(1, 6 * mm))

        # 3. 词汇等级 pill
        vocab_level = _compute_vocab_level(items)
        if vocab_level != '—':
            pill_text = f'词汇等级  {vocab_level}'
            pill = Paragraph(
                f'<font color="#FFFFFF"><b>  {pill_text}  </b></font>',
                ParagraphStyle('cover_pill', parent=styles['cover_stat'],
                               fontSize=11, leading=16, alignment=TA_CENTER,
                               backColor=BRAND_GREEN, borderPadding=4)
            )
            flow.append(pill)

    flow.append(Spacer(1, 8 * mm))

    # 统计: 单词 N · 短语 N · 表达 N
    stat_parts = []
    if stats['words']:
        stat_parts.append(f"单词 <b>{stats['words']}</b>")
    if stats['phrases']:
        stat_parts.append(f"短语 <b>{stats['phrases']}</b>")
    if stats['expressions']:
        stat_parts.append(f"表达 <b>{stats['expressions']}</b>")
    if stat_parts:
        flow.append(Paragraph(' · '.join(stat_parts), styles['cover_stat']))
        flow.append(Spacer(1, 4 * mm))
        total = stats['words'] + stats['phrases'] + stats['expressions']
        flow.append(Paragraph(f'共 {total} 条学习内容', styles['cover_stat']))

    flow.append(PageBreak())
    return flow


def _build_toc_page(stats: dict, styles: dict) -> list:
    """PDF 版 TOC 目录页 (Phase B)

    视觉: 居中标题 "📖 目录" + 3 行 entry (section 名 + count)
    PDF 不做精确页码 (需要 2-pass, 复杂度高), 只列 section + count,
    用户可以靠页眉 / 滚动定位
    """
    flow = []
    flow.append(Spacer(1, 50 * mm))

    # 标题 (用品牌色, 跟 DOCX 一致)
    # 不用 emoji: reportlab 注册的 NotoSansSC / DejaVuSans 都不含 Plane 1+ 字符,
    # 📖 会渲染成 □. 用纯文字 + letter-spacing 视觉替代
    flow.append(Paragraph('目  录', ParagraphStyle(
        'toc_title', parent=styles['h1_cover'],
        fontSize=28, textColor=BRAND_GREEN, alignment=TA_CENTER,
    )))
    flow.append(Spacer(1, 4 * mm))
    flow.append(Paragraph('<i>Table of Contents</i>', ParagraphStyle(
        'toc_sub', parent=styles['h2_cover_sub'],
        fontSize=10, textColor=TEXT_MUTED, alignment=TA_CENTER,
    )))
    flow.append(Spacer(1, 25 * mm))

    # 3 个 entry (section 名 + count, 居中显示, 简洁)
    entries = []
    if stats.get('words'):
        entries.append(('单词卡', stats['words']))
    if stats.get('phrases'):
        entries.append(('短语卡', stats['phrases']))
    if stats.get('expressions'):
        entries.append(('表达卡', stats['expressions']))

    toc_style = ParagraphStyle(
        'toc_entry', parent=styles['h2_cover_sub'],
        fontSize=15, leading=32, textColor=TEXT_PRIMARY,
        alignment=TA_CENTER,
    )
    for name, count in entries:
        # 用 ◦ 圆点 + 全角空格装饰, 视觉对齐
        flow.append(Paragraph(
            f'<font color="#4DA06C"><b>{name}</b></font>'
            f'<font color="#94A3B8">  ·  </font>'
            f'<font color="#64748B">{count} 个</font>',
            toc_style,
        ))

    flow.append(PageBreak())
    return flow


def _build_section_banner(title, count, styles):
    """板块标题 banner (单词卡 / 短语卡 / 表达卡)

    Phase D: 在 banner 后追加一行章节导言, 让用户对板块有 1 句话预期
    """
    txt = f'{title} <font size="11" color="#94A3B8">({count})</font>'
    p = Paragraph(txt, styles['section_banner'])
    # banner 用一个带背景色的单 cell table
    t = Table([[p]], colWidths=[OUTPUT_PAGE_SIZE[0] - PAGE_MARGIN_LEFT - PAGE_MARGIN_RIGHT])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), BRAND_GREEN_SUBTLE),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING', (0, 0), (-1, -1), 12),
        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ('LINEBELOW', (0, 0), (-1, -1), 2, BRAND_GREEN),
    ]))

    flow = [t]
    # Phase D: 章节导言 (1 句)
    intro_map = {
        '单词卡': f'本节收录 <b>{count}</b> 个视频重点单词, 每条含音标 / 释义 / 例句, 建议先听后读。',
        '短语卡': f'本节收录 <b>{count}</b> 个常用短语搭配, 注意短语整体记忆而非拆字。',
        '表达卡': f'本节收录 <b>{count}</b> 个高级表达, 含结构解析 / 举一反三 / 适用场景, 可作为口语 / 写作素材。',
    }
    intro_text = intro_map.get(title)
    if intro_text:
        intro_p = Paragraph(intro_text, ParagraphStyle(
            'section_intro', parent=styles['meta_text'],
            fontSize=10, leading=14, textColor=TEXT_SECONDARY,
            leftIndent=4, rightIndent=4, spaceBefore=4, spaceAfter=8,
        ))
        flow.append(intro_p)

    flow.append(Spacer(1, 3 * mm))
    return flow


def _build_words_or_phrases_table(items, fields, styles, start_no=1):
    # 表头
    header = [
        Paragraph('No.', styles['table_header']),
        Paragraph('Word', styles['table_header']),
        Paragraph('Meaning', styles['table_header']),
    ]
    table_data = [header]

    for idx, item in enumerate(items, start=start_no):
        # === Word 列: 英文 + 音标 + (完整版) 难度/词频 ===
        word_parts = []
        word_parts.append(Paragraph(_escape_xml(item.content_en), styles['word_text']))
        phonetic = _safe_text(getattr(item, 'phonetic', None))
        if phonetic:
            # 音标段落单独用 IPA 字体, 不能用主字体 (YaHei 缺 ˈ ɜː ɪ ʃ 等字符)
            # DB 里 phonetic 通常带两边的 /, 需要先剥掉再统一加一层 /,
            # 避免出现 //..// 双重斜杠
            p = phonetic.strip().strip('/').strip()
            word_parts.append(Paragraph(_escape_xml(f'/{p}/'), styles['phonetic_ipa']))

        if fields == 'full':
            # 难度 badge
            diff = getattr(item, 'difficulty', None)
            if diff and 1 <= diff <= 5:
                diff_color = DIFFICULTY_COLORS[diff - 1].hexval()[2:]  # 去 0x
                word_parts.append(Paragraph(
                    f'<font color="#{diff_color}">●</font> '
                    f'<font size="9" color="#94A3B8">难度 {diff}/5</font>',
                    styles['phonetic']))
            freq = getattr(item, 'frequency_rank', None)
            if freq:
                word_parts.append(Paragraph(
                    f'<font size="9" color="#94A3B8">词频 Top {freq}</font>',
                    styles['phonetic']))

        # === Meaning 列: 词性+中文 + 例句 + (完整版) 其他 ===
        meaning_parts = []
        pos = _safe_text(getattr(item, 'part_of_speech', None))
        cn = _safe_text(getattr(item, 'content_cn', None))
        # content_cn 通常已经包含词性 (如 "n. 家具"), 避免重复拼接 → "n. n. 家具"
        if pos and cn:
            # 比较 pos 去掉末尾句号后的形式 ("n." -> "n")
            pos_clean = pos.lower().rstrip('.').strip()
            cn_lstrip = cn.lstrip()
            # cn 第一个 token ("n." "v." "adj." 等) 是否匹配 pos
            cn_first = cn_lstrip.split()[0].lower().rstrip('.').strip() if cn_lstrip else ''
            if cn_first == pos_clean:
                head = _escape_xml(cn)
            else:
                head = f'<b>{_escape_xml(pos)}</b> {_escape_xml(cn)}'
        elif pos:
            head = f'<b>{_escape_xml(pos)}</b>'
        elif cn:
            head = _escape_xml(cn)
        else:
            head = ''
        if head:
            meaning_parts.append(Paragraph(head, styles['pos_cn']))

        if fields == 'full':
            en_def = _safe_text(getattr(item, 'english_definition', None))
            if en_def:
                meaning_parts.append(Paragraph(
                    f'<font size="9" color="#94A3B8">英英: {_escape_xml(en_def)}</font>',
                    styles['example_cn']))
            other_pos = _safe_text(getattr(item, 'other_pos_definitions', None))
            if other_pos:
                meaning_parts.append(Paragraph(
                    f'<font size="9" color="#94A3B8">其他词性: {_escape_xml(other_pos)}</font>',
                    styles['example_cn']))

        # 例句 (优先 example_sentence, 没有就 fallback context_sentence)
        example_en = _safe_text(getattr(item, 'example_sentence', None))
        if not example_en:
            example_en = _safe_text(getattr(item, 'context_sentence', None))
        if example_en:
            meaning_parts.append(Paragraph(
                f'<font color="#4DA06C"><b>EN</b></font>  {_escape_xml(example_en)}',
                styles['example_en']))
            # 字幕出处中文翻译
            ctx_cn = _safe_text(getattr(item, 'context_translation', None))
            if ctx_cn:
                meaning_parts.append(Paragraph(
                    f'<font color="#94A3B8"><b>CN</b></font>  {_escape_xml(ctx_cn)}',
                    styles['example_cn']))

        # No. 列
        no_p = Paragraph(str(idx), styles['no'])

        # 行 cell — word 列垂直拼接用 nested table 或直接 list
        # reportlab Table cell 可以接受 list of flowables, 自动垂直排列
        table_data.append([no_p, word_parts, meaning_parts])

    # 列宽: No. 窄 / Word 中 / Meaning 宽
    avail_width = OUTPUT_PAGE_SIZE[0] - PAGE_MARGIN_LEFT - PAGE_MARGIN_RIGHT
    col_widths = [avail_width * 0.06, avail_width * 0.32, avail_width * 0.62]

    t = Table(table_data, colWidths=col_widths, repeatRows=1)
    style_cmds = [
        # 表头
        ('BACKGROUND', (0, 0), (-1, 0), BRAND_GREEN),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), _FONT_BOLD),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('ALIGN', (0, 0), (0, -1), 'CENTER'),     # No. 居中
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        # 行交替 (跳过表头)
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [BRAND_GREEN_LIGHT, colors.white]),
        # 框线
        ('LINEBELOW', (0, 0), (-1, 0), 0, colors.transparent),
        ('LINEBELOW', (0, 1), (-1, -1), 0.3, BORDER_COLOR),
        ('BOX', (0, 0), (-1, -1), 0.6, BRAND_GREEN),
    ]
    t.setStyle(TableStyle(style_cmds))
    return t


def _build_expression_card(item, idx, styles):
    """单个表达卡 (独立色块)"""
    bg_color = ACCENT_PURPLE if idx % 2 == 0 else ACCENT_PINK

    parts = []
    # 原句 (英文) - 大字标题
    parts.append(Paragraph(_escape_xml(item.content_en), styles['expression_title']))
    # 中文释义 (加 前缀 区分)
    cn = _safe_text(getattr(item, 'content_cn', None))
    if cn:
        parts.append(Paragraph(
            f'<font color="#94A3B8"><b>中文</b></font>  '
            f'<font color="#475569">{_escape_xml(cn)}</font>',
            styles['phonetic']))
    parts.append(Spacer(1, 2 * mm))

    # 结构解析 / 举一反三 / 使用场景 / 相似表达
    def _add_line(label, value, icon=''):
        if value:
            # Phase F3: PDF 不直接用 emoji (NotoSansSC / DejaVuSans 都不含 Plane 1+)
            # 用 ● 圆点 + label 替代, 跟表达卡视觉一致
            icon_str = f'<font color="#4DA06C">●</font>  ' if icon else ''
            parts.append(Paragraph(
                f'{icon_str}'
                f'<font color="#4DA06C"><b>{label}</b></font>  '
                f'<font color="#1E293B">{_escape_xml(value)}</font>',
                styles['example_en']))

    _add_line('结构', _safe_text(getattr(item, 'structure_analysis', None)), icon='1')
    _add_line('举一反三', _safe_text(getattr(item, 'similar_expressions', None)), icon='1')
    _add_line('场景', _safe_text(getattr(item, 'usage_scenario', None)), icon='1')
    _add_line('相似表达', _safe_text(getattr(item, 'alternative_phrasings', None)), icon='1')

    # 例句
    example_en = _safe_text(getattr(item, 'example_sentence', None))
    if example_en:
        _add_line('例句', example_en, icon='1')
    ctx_cn = _safe_text(getattr(item, 'context_translation', None))
    if ctx_cn:
        parts.append(Paragraph(
            f'<font color="#94A3B8"><b>CN</b></font>  {_escape_xml(ctx_cn)}',
            styles['example_cn']))

    # 用单 cell table 套色块背景
    avail_width = OUTPUT_PAGE_SIZE[0] - PAGE_MARGIN_LEFT - PAGE_MARGIN_RIGHT
    card = Table([[parts]], colWidths=[avail_width])
    card.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), bg_color),
        ('LEFTPADDING', (0, 0), (-1, -1), 12),
        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('LINEBEFORE', (0, 0), (0, -1), 3, BRAND_GREEN),  # 左侧绿色竖线
    ]))
    return [KeepTogether([card, Spacer(1, 4 * mm)])]


def export_pdf(material_title: str, interpretations: Iterable, fields: str = 'compact') -> bytes:
    """生成学习卡片 PDF

    Args:
        material_title: 视频标题
        interpretations: VideoInterpretation 列表
        fields: 'compact' 或 'full'

    Returns:
        PDF 文件 bytes
    """
    material_title = _sanitize_title(material_title)
    _ensure_font_registered()
    styles = _make_styles()

    groups = _categorize(interpretations)
    stats = {k: len(v) for k, v in groups.items()}
    today_str = datetime.date.today().strftime('%Y-%m-%d')

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=OUTPUT_PAGE_SIZE,
        topMargin=PAGE_MARGIN_TOP, bottomMargin=PAGE_MARGIN_BOTTOM,
        leftMargin=PAGE_MARGIN_LEFT, rightMargin=PAGE_MARGIN_RIGHT,
        title=f'{material_title} - 学习卡',
        author='Linyu',
    )

    story = []
    # 1. 封面
    story.extend(_build_cover(material_title, stats, today_str, styles, items=interpretations))

    # 2. Phase B: TOC 目录页
    if stats['words'] or stats['phrases'] or stats['expressions']:
        story.extend(_build_toc_page(stats, styles))

    # 2. 单词卡
    if groups['words']:
        story.extend(_build_section_banner('单词卡', stats['words'], styles))
        story.append(_build_words_or_phrases_table(groups['words'], fields, styles, start_no=1))
        story.append(Spacer(1, 10 * mm))

    # 3. 短语卡
    if groups['phrases']:
        # CondPageBreak: 只在当前页剩余空间不够 30mm 时才换页,
        # 避免短语卡刚塞满一页又强制 PageBreak 留个空白页
        story.append(CondPageBreak(30 * mm))
        story.extend(_build_section_banner('短语卡', stats['phrases'], styles))
        start = stats['words'] + 1  # 编号延续? 还是每个板块独立从 1 开始?
        # 参考版是独立编号, 我们也独立从 1 开始更直观
        story.append(_build_words_or_phrases_table(groups['phrases'], fields, styles, start_no=1))
        story.append(Spacer(1, 10 * mm))

    # 4. 表达卡
    if groups['expressions']:
        story.append(CondPageBreak(30 * mm))
        story.extend(_build_section_banner('表达卡', stats['expressions'], styles))
        for idx, item in enumerate(groups['expressions']):
            story.extend(_build_expression_card(item, idx, styles))

    # 渲染
    doc.build(
        story,
        onFirstPage=lambda c, d: _on_page(c, d, material_title),
        onLaterPages=lambda c, d: _on_page(c, d, material_title),
    )
    return buf.getvalue()


# ==================== Word (.docx) 构建 ====================

def _set_cell_shading(cell, hex_color: str):
    """给 docx 单元格加背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color='E2E8F0', size='4'):
    """单元格四边边框"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _set_cell_borders_none(cell):
    """清空单元格所有边框 (用于装饰条 / 色块不想显示边框的场景)"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'nil')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _set_cell_margins(cell, top_pt=8, bottom_pt=8, left_pt=8, right_pt=8):
    """设置单元格内边距 (tcMar, 单位 pt)"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in (('top', top_pt), ('left', left_pt),
                      ('bottom', bottom_pt), ('right', right_pt)):
        node = OxmlElement(f'w:{edge}')
        # 单位: 20 twips = 1 pt → val pt = val * 20 twips
        node.set(qn('w:w'), str(int(val * 20)))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tc_pr.append(tcMar)


def _set_paragraph_bottom_border(paragraph, size_eighths=24, color='4DA06C', space_pt=4):
    """段落加底部边框 (八分之一磅为单位, 24 = 3pt)

    对标 PDF 的 LINEBELOW / LINE 装饰线
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    # 避免重复加
    existing = pBdr.find(qn('w:bottom'))
    if existing is not None:
        pBdr.remove(existing)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size_eighths))
    bottom.set(qn('w:space'), str(int(space_pt * 20)))  # space 单位 twips
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)


def _add_horizontal_decoration(doc, width_mm=60, color='4DA06C'):
    """封面装饰横线: 居中的 Nmm 宽绿条 (用 1×3 table 模拟, 中间 cell 着色, 两侧 transparent)"""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    page_w_mm = 210  # A4
    margin_mm = 16
    avail_mm = page_w_mm - 2 * margin_mm
    side_mm = (avail_mm - width_mm) / 2
    tbl.columns[0].width = Mm(side_mm)
    tbl.columns[1].width = Mm(width_mm)
    tbl.columns[2].width = Mm(side_mm)
    # 中间 cell 着色
    middle = tbl.cell(0, 1)
    _set_cell_shading(middle, color)
    _set_cell_borders_none(middle)
    # 两侧 cell 透明 + 无边框 + 行高尽量小 (0)
    for cell in (tbl.cell(0, 0), tbl.cell(0, 2)):
        _set_cell_borders_none(cell)
        _set_cell_shading(cell, 'FFFFFF')  # 显式白底避免主题色透出
        # 段落置空 (避免默认空段落占空间)
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
    return tbl


def _add_para_run(paragraph, text, *, bold=False, italic=False, size=11,
                  color_hex=None, font_name=None):
    """给 paragraph 加一个 run, 自动设中英文字体"""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    # 中英文字体
    name = font_name or 'Microsoft YaHei'
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    return run


# ==================== DOCX 书签 + 内部超链接 (Phase B: TOC 跳转用) ====================

_bookmark_counter = [1000]  # 模块级单调 id, 每个 export 调用开头 reset


def _reset_bookmark_counter():
    _bookmark_counter[0] = 1000


def _next_bookmark_id() -> int:
    _bookmark_counter[0] += 1
    return _bookmark_counter[0]


# 固定的 bookmark name, 全文档唯一 (Phase B)
BOOKMARK_WORDS = 'sec_words'
BOOKMARK_PHRASES = 'sec_phrases'
BOOKMARK_EXPRESSIONS = 'sec_expressions'


def _add_bookmark_around_paragraph(paragraph, bookmark_name: str):
    """段落首尾包 bookmarkStart/bookmarkEnd, 供后续 _add_internal_hyperlink 跳转"""
    bid = _next_bookmark_id()
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bid))
    start.set(qn('w:name'), bookmark_name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bid))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_internal_hyperlink(paragraph, text: str, bookmark_name: str,
                            *, size=11, color_hex='4DA06C', bold=False) -> None:
    """段落里插一个内部超链接 (指向 bookmark), 默认下划线 + 品牌绿

    OXML 结构:
      <w:hyperlink w:anchor="bookmark_name">
        <w:r>
          <w:rPr><w:rFonts/><w:color/><w:sz/><w:u val=single/></w:rPr>
          <w:t>text</w:t>
        </w:r>
      </w:hyperlink>
    """
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    hyperlink.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 字体三槽位都走 Microsoft YaHei (含中文 fallback)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
    rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
    rPr.append(rFonts)

    if color_hex:
        color = OxmlElement('w:color')
        color.set(qn('w:val'), color_hex)
        rPr.append(color)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))  # 半磅
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size * 2)))
    rPr.append(szCs)

    if bold:
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ==================== IPA 字体 (Phase F1) ====================

# Microsoft YaHei 缺 IPA 字符 (ˈ ɜː ɪ ʃ ŋ θ ð ...), 改用 DejaVu Sans (跨平台自带, 含完整 IPA)
_DOCX_IPA_FONT = 'DejaVu Sans'


def _set_run_ipa_font(run):
    """把 run 的 ascii/hAnsi/cs 槽位设成 DejaVu Sans, eastAsia 保留 YaHei

    注意: 必须配合 _add_para_run 之前/之后调用, 因为 _add_para_run 会强制覆盖全部三槽
    → 调用顺序: 先 _add_para_run (默认 YaHei), 再 _set_run_ipa_font (覆盖 ascii/hAnsi/cs)
    """
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), _DOCX_IPA_FONT)
    rFonts.set(qn('w:hAnsi'), _DOCX_IPA_FONT)
    rFonts.set(qn('w:cs'), _DOCX_IPA_FONT)
    # eastAsia 不动, 中文仍 YaHei


# ==================== DOCX TOC 目录页 (Phase B) ====================

def _add_index_page(doc, words: list) -> None:
    """末尾页: A-Z 词汇索引
    - 按首字母 (A-Z) 分组, 非字母开头归到 "#"
    - 每个 word 是 hyperlink → BOOKMARK_WORDS (跳转后用户用 Word 查找定位具体词)
    - 字母 heading 走品牌绿 + 底部分隔线
    - 每个 letter 下 word 列表, 每行 ~6 个用 · 分隔, 紧凑排版
    """
    if not words:
        return

    doc.add_page_break()

    # 顶部空行把标题推到 1/3 处
    for _ in range(3):
        doc.add_paragraph('')

    # 标题 (emoji 由 Word 用 Segoe UI Emoji fallback 渲染, 跟封面 ✅ 一致)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_para_run(p, '🔤  词汇索引', bold=True, size=22, color_hex='4DA06C')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_para_run(p, 'Vocabulary Index', italic=True, size=10, color_hex='94A3B8')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    _add_para_run(p, f'共 {len(words)} 个词条  ·  按首字母排序  ·  点击跳转',
                  size=10, color_hex='94A3B8')

    for _ in range(2):
        doc.add_paragraph('')

    # 按首字母分组
    groups_dict = {}
    for w in words:
        if not w:
            continue
        first = w[0].upper()
        if not ('A' <= first <= 'Z'):
            first = '#'
        groups_dict.setdefault(first, []).append(w)

    # 字母表顺序, '#' 放最后
    sorted_letters = sorted(groups_dict.keys(), key=lambda c: (c == '#', c))

    for letter in sorted_letters:
        # 字母 heading
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(10)
        ph.paragraph_format.space_after = Pt(4)
        ph.paragraph_format.left_indent = Cm(0.4)
        _add_para_run(ph, letter, bold=True, size=20, color_hex='4DA06C')
        # 底部分隔线 (对标 PDF 章节标题 2pt 绿条)
        _set_paragraph_bottom_border(ph, size_eighths=12, color='4DA06C', space_pt=2)

        # 该字母下的 words (按字母序, 忽略大小写)
        word_list = sorted(groups_dict[letter], key=str.lower)
        # 每行 ~6 个, 用 · 分隔
        for chunk_start in range(0, len(word_list), 6):
            chunk = word_list[chunk_start:chunk_start + 6]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(3)
            for i, w in enumerate(chunk):
                if i > 0:
                    _add_para_run(p, ' · ', size=11, color_hex='CBD5E1')
                # hyperlink → BOOKMARK_WORDS
                _add_internal_hyperlink(p, w, BOOKMARK_WORDS,
                                        size=11, color_hex='1E293B')

        # 字母块之间用空段分隔
        doc.add_paragraph('')


def _add_toc_page(doc, stats: dict) -> None:
    """在封面后插入 TOC 目录页, 3 个 section 各一行 entry

    视觉: 居中标题 "📖 目录" + 表格 (左=section 名 hyperlink, 中=虚线, 右=count)
    点击左 entry 跳转到对应板块 (靠 _add_section_banner 注入的 bookmark)
    """
    # 顶部空行把标题推到页面 1/3
    for _ in range(4):
        doc.add_paragraph('')

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_para_run(p, '📖  目录', bold=True, size=22, color_hex='4DA06C')

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_para_run(p, 'Table of Contents', italic=True, size=10, color_hex='94A3B8')

    # 推开
    for _ in range(3):
        doc.add_paragraph('')

    # 3 行 entry (用 1×3 表格模拟 "............" 视觉)
    entries = []
    if stats.get('words'):
        entries.append(('单词卡', stats['words'], BOOKMARK_WORDS))
    if stats.get('phrases'):
        entries.append(('短语卡', stats['phrases'], BOOKMARK_PHRASES))
    if stats.get('expressions'):
        entries.append(('表达卡', stats['expressions'], BOOKMARK_EXPRESSIONS))

    for name, count, bookmark in entries:
        # 用 1×3 表格: 左 (section 名 hyperlink) / 中 (虚线) / 右 (count)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        col_w = (Mm(40), Mm(80), Mm(40))
        for i, w in enumerate(col_w):
            for cell in tbl.columns[i].cells:
                cell.width = w
        left, mid, right = tbl.rows[0].cells

        # 三个 cell 都去掉边框
        for c in (left, mid, right):
            _set_cell_borders_none(c)

        # 左: section 名 (hyperlink → bookmark)
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_internal_hyperlink(p, f'{name} ', bookmark, size=14, color_hex='1E293B', bold=True)
        _add_para_run(p, '  · ', size=14, color_hex='94A3B8')

        # 中: 虚线占位 (用重复 · 字符)
        p = mid.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_para_run(p, '· ' * 25, size=12, color_hex='CBD5E1')

        # 右: count + 个
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _add_para_run(p, f'  {count} 个', size=14, color_hex='64748B')

        # 行间距
        for c in (left, mid, right):
            for para in c.paragraphs:
                para.paragraph_format.space_before = Pt(8)
                para.paragraph_format.space_after = Pt(8)

    doc.add_page_break()


def export_docx(material_title: str, interpretations: Iterable, fields: str = 'compact') -> bytes:
    """生成学习卡片 Word 文档"""
    material_title = _sanitize_title(material_title)
    doc = Document()
    # Phase B: 重置 bookmark 计数器, 保证本次 export 内 id 唯一
    _reset_bookmark_counter()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(18)
        section.left_margin = Mm(16)
        section.right_margin = Mm(16)

    groups = _categorize(interpretations)
    stats = {k: len(v) for k, v in groups.items()}
    today_str = datetime.date.today().strftime('%Y-%m-%d')

    # === 封面 ===
    # 用空行推到页面中部
    for _ in range(6):
        doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_para_run(p, 'Linyu 学习卡', size=12, color_hex='64748B')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_para_run(p, material_title or '视频学习卡片',
                  bold=True, size=26, color_hex='4DA06C')
    # 装饰横线 (对标 PDF 60mm 绿条)
    _add_horizontal_decoration(doc, width_mm=60, color='4DA06C')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_para_run(p, today_str, size=12, color_hex='64748B')

    # ===== Phase C: 3 项数据可视化 (跟 PDF 对齐) =====
    # 1. 阅读时长
    mins = _compute_reading_minutes(
        stats['words'], stats['phrases'], stats['expressions'])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    _add_para_run(p, '预计学习时长  ', size=12, color_hex='94A3B8')
    _add_para_run(p, f'≈ {mins} 分钟', bold=True, size=14, color_hex='1E293B')

    # 2. 难度分布条 (1×N table, 各 cell 按比例着色)
    items = list(interpretations) if not isinstance(interpretations, list) else interpretations
    diff_dist = _compute_difficulty_distribution(items)
    total_diff = sum(diff_dist.values())
    if total_diff > 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        bar_tbl = doc.add_table(rows=1, cols=5)
        bar_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        bar_tbl.autofit = False
        # 总宽 ~ 90mm, 按比例分配, 最少 8mm (保证可见)
        avail_mm = 90
        cells = bar_tbl.rows[0].cells
        for d in range(1, 6):
            n = diff_dist[d]
            w_mm = max(8, avail_mm * n / total_diff)
            cells[d - 1].width = Mm(w_mm)
            _set_cell_shading(cells[d - 1], DIFFICULTY_COLORS_HEX[d - 1])
            _set_cell_margins(cells[d - 1], top_pt=2, bottom_pt=2, left_pt=0, right_pt=0)
            _set_cell_borders_none(cells[d - 1])
        # 难度图例 (简单 ●●●●● 困难)
        legend = doc.add_paragraph()
        legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
        legend.paragraph_format.space_before = Pt(4)
        _add_para_run(legend, '简单 ', size=10, color_hex='94A3B8')
        for d in range(1, 6):
            _add_para_run(legend, '●', size=12, color_hex=DIFFICULTY_COLORS_HEX[d - 1])
        _add_para_run(legend, ' 困难', size=10, color_hex='94A3B8')

    # 3. 词汇等级 pill (1×1 表格 + 居中文字)
    vocab_level = _compute_vocab_level(items)
    if vocab_level != '—':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        pill_tbl = doc.add_table(rows=1, cols=1)
        pill_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        pill_cell = pill_tbl.rows[0].cells[0]
        _set_cell_shading(pill_cell, '4DA06C')
        _set_cell_borders_none(pill_cell)
        _set_cell_margins(pill_cell, top_pt=6, bottom_pt=6, left_pt=20, right_pt=20)
        pp = pill_cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_para_run(pp, f'词汇等级  ', bold=True, size=12, color_hex='FFFFFF')
        _add_para_run(pp, vocab_level, bold=True, size=12, color_hex='FFFFFF')

    # 装饰线与统计之间推开
    for _ in range(2):
        doc.add_paragraph('')
    # 统计
    stat_parts = []
    if stats['words']:
        stat_parts.append(f'单词 {stats["words"]}')
    if stats['phrases']:
        stat_parts.append(f'短语 {stats["phrases"]}')
    if stats['expressions']:
        stat_parts.append(f'表达 {stats["expressions"]}')
    if stat_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_para_run(p, ' · '.join(stat_parts), size=12, color_hex='64748B')

    doc.add_page_break()

    # === Phase B: TOC 目录页 ===
    # 封面后, 板块标题前
    # 用 1×3 表格 (左 entry / 中虚线 / 右 entry) 模拟 "............" 视觉
    # 实际是 hyperlink, 用户在 Word 里点会跳到对应板块
    _add_toc_page(doc, stats)

    # === 板块标题 helper (Phase B: 加 bookmark, 让 TOC 可跳转; Phase D: 加章节导言) ===
    def _add_section_banner(title, count, bookmark_name=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_para_run(p, title, bold=True, size=18, color_hex='4DA06C')
        _add_para_run(p, f'  ({count})', size=11, color_hex='94A3B8')
        # banner 段落底色: 给 paragraph 加 shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), 'E6F4EC')
        pPr.append(shd)
        # 上下 padding (对标 PDF TOPPADDING/BOTTOMPADDING 5pt)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        # 底部分隔线 (对标 PDF LINEBELOW 2pt 绿)
        _set_paragraph_bottom_border(p, size_eighths=16, color='4DA06C', space_pt=4)
        # Phase B: 段首加 bookmark (供 TOC 跳转)
        if bookmark_name:
            _add_bookmark_around_paragraph(p, bookmark_name)
        # Phase D: 章节导言 (1 句)
        intro_map = {
            '单词卡': f'本节收录 {count} 个视频重点单词, 每条含音标 / 释义 / 例句, 建议先听后读。',
            '短语卡': f'本节收录 {count} 个常用短语搭配, 注意短语整体记忆而非拆字。',
            '表达卡': f'本节收录 {count} 个高级表达, 含结构解析 / 举一反三 / 适用场景, 可作为口语 / 写作素材。',
        }
        intro_text = intro_map.get(title)
        if intro_text:
            ip = doc.add_paragraph()
            ip.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ip.paragraph_format.left_indent = Cm(0.4)
            ip.paragraph_format.right_indent = Cm(0.4)
            ip.paragraph_format.space_before = Pt(4)
            ip.paragraph_format.space_after = Pt(8)
            _add_para_run(ip, intro_text, size=10, color_hex='64748B', italic=True)

    def _add_words_table(items):
        if not items:
            return
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # 表头
        hdr = table.rows[0].cells
        headers = ['No.', 'Word', 'Meaning']
        for i, h in enumerate(headers):
            _set_cell_shading(hdr[i], '4DA06C')
            _set_cell_borders(hdr[i], color='4DA06C', size='4')
            p = hdr[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            _add_para_run(p, h, bold=True, size=11, color_hex='FFFFFF')

        for idx, item in enumerate(items, start=1):
            row = table.add_row().cells
            # 行交替底色
            if idx % 2 == 1:
                for c in row:
                    _set_cell_shading(c, 'F0F9F3')
            for c in row:
                _set_cell_borders(c, color='E2E8F0', size='4')

            # No.
            p0 = row[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_para_run(p0, str(idx), size=10, color_hex='94A3B8')

            # Word 列
            p1 = row[1].paragraphs[0]
            _add_para_run(p1, _safe_text(item.content_en), bold=True, size=12, color_hex='4DA06C')
            phonetic = _safe_text(getattr(item, 'phonetic', None))
            if phonetic:
                p1b = row[1].add_paragraph()
                # DB 里 phonetic 通常带两边的 /, 先剥再统一加一层避免 //..//
                p = phonetic.strip().strip('/').strip()
                run = _add_para_run(p1b, f'/{p}/', italic=True, size=10, color_hex='64748B')
                # Phase F1: 把 IPA 字符 (ˈ ɜː ɪ ʃ ŋ θ ð ...) 切到 DejaVu Sans,
                # 避免在 Microsoft YaHei 下显示成方框
                _set_run_ipa_font(run)
            if fields == 'full':
                diff = getattr(item, 'difficulty', None)
                if diff and 1 <= diff <= 5:
                    p1c = row[1].add_paragraph()
                    # Phase F2: 难度彩色圆点 + 文字 (跟 PDF 对齐)
                    # N 个 ● 每个不同 hex 色, 颜色顺序按 DIFFICULTY_COLORS_HEX
                    for d_idx in range(diff):
                        _add_para_run(p1c, '●', size=11,
                                      color_hex=DIFFICULTY_COLORS_HEX[d_idx])
                    _add_para_run(p1c, f'  难度 {diff}/5', size=9, color_hex='94A3B8')
                freq = getattr(item, 'frequency_rank', None)
                if freq:
                    p1d = row[1].add_paragraph()
                    _add_para_run(p1d, f'词频 Top {freq}', size=9, color_hex='94A3B8')

            # Meaning 列
            p2 = row[2].paragraphs[0]
            pos = _safe_text(getattr(item, 'part_of_speech', None))
            cn = _safe_text(getattr(item, 'content_cn', None))
            # 跟 PDF 一致: content_cn 通常已含词性 ("n. 家具"), 避免重复拼
            cn_already_has_pos = False
            if pos and cn:
                pos_clean = pos.lower().rstrip('.').strip()
                cn_first = cn.lstrip().split()[0].lower().rstrip('.').strip() if cn.lstrip() else ''
                cn_already_has_pos = cn_first == pos_clean
            if pos and not cn_already_has_pos:
                _add_para_run(p2, pos + ' ', bold=True, size=11, color_hex='1E293B')
            if cn:
                _add_para_run(p2, cn, size=11, color_hex='1E293B')

            if fields == 'full':
                en_def = _safe_text(getattr(item, 'english_definition', None))
                if en_def:
                    p2b = row[2].add_paragraph()
                    _add_para_run(p2b, '英英: ', size=9, color_hex='94A3B8')
                    _add_para_run(p2b, en_def, size=9, color_hex='94A3B8')
                other_pos = _safe_text(getattr(item, 'other_pos_definitions', None))
                if other_pos:
                    p2c = row[2].add_paragraph()
                    _add_para_run(p2c, '其他词性: ', size=9, color_hex='94A3B8')
                    _add_para_run(p2c, other_pos, size=9, color_hex='94A3B8')

            # 例句
            example_en = _safe_text(getattr(item, 'example_sentence', None))
            if not example_en:
                example_en = _safe_text(getattr(item, 'context_sentence', None))
            if example_en:
                p2d = row[2].add_paragraph()
                _add_para_run(p2d, 'EN  ', bold=True, size=10, color_hex='4DA06C')
                _add_para_run(p2d, example_en, size=10, color_hex='1E293B')
                ctx_cn = _safe_text(getattr(item, 'context_translation', None))
                if ctx_cn:
                    p2e = row[2].add_paragraph()
                    _add_para_run(p2e, 'CN  ', bold=True, size=10, color_hex='94A3B8')
                    _add_para_run(p2e, ctx_cn, size=10, color_hex='64748B')

        # 列宽 (对标 PDF: 6% / 32% / 62% × 178mm 可用宽度 ≈ 11 / 57 / 110mm)
        for row in table.rows:
            row.cells[0].width = Cm(1.1)
            row.cells[1].width = Cm(5.7)
            row.cells[2].width = Cm(11.0)

    # === 单词 ===
    if groups['words']:
        _add_section_banner('单词卡', stats['words'], bookmark_name=BOOKMARK_WORDS)
        _add_words_table(groups['words'])
        doc.add_paragraph('')

    # === 短语 ===
    if groups['phrases']:
        doc.add_page_break()
        _add_section_banner('短语卡', stats['phrases'], bookmark_name=BOOKMARK_PHRASES)
        _add_words_table(groups['phrases'])
        doc.add_paragraph('')

    # === 表达 ===
    if groups['expressions']:
        doc.add_page_break()
        _add_section_banner('表达卡', stats['expressions'], bookmark_name=BOOKMARK_EXPRESSIONS)
        for idx, item in enumerate(groups['expressions']):
            bg = 'F5F0FF' if idx % 2 == 0 else 'FFF0F5'
            # 每个表达用单 cell table 当色块
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.rows[0].cells[0]
            _set_cell_shading(cell, bg)
            _set_cell_borders(cell, color='E2E8F0', size='4')
            # 对标 PDF: 8pt 四边 padding, 给内容呼吸空间
            _set_cell_margins(cell, top_pt=8, bottom_pt=8, left_pt=12, right_pt=12)
            # 左侧绿色竖线 (通过 left border 加粗加色)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn('w:tcBorders'))
            if tc_borders is None:
                tc_borders = OxmlElement('w:tcBorders')
                tc_pr.append(tc_borders)
            # 替换 left border
            existing_left = tc_borders.find(qn('w:left'))
            if existing_left is not None:
                tc_borders.remove(existing_left)
            left_b = OxmlElement('w:left')
            left_b.set(qn('w:val'), 'single')
            left_b.set(qn('w:sz'), '24')   # 3pt
            left_b.set(qn('w:color'), '4DA06C')
            tc_borders.append(left_b)

            # 内容
            p = cell.paragraphs[0]
            _add_para_run(p, _safe_text(item.content_en), bold=True, size=12, color_hex='1E293B')
            cn = _safe_text(getattr(item, 'content_cn', None))
            if cn:
                p2 = cell.add_paragraph()
                _add_para_run(p2, cn, size=10, color_hex='94A3B8')

            def _add_field(label, val, icon=''):
                if val:
                    pf = cell.add_paragraph()
                    if icon:
                        _add_para_run(pf, f'{icon} ', bold=True, size=11, color_hex='4DA06C')
                    _add_para_run(pf, f'{label}  ', bold=True, size=10, color_hex='4DA06C')
                    _add_para_run(pf, val, size=10, color_hex='1E293B')

            # Phase F3: 表达卡字段加 icon (跟参考 PDF 对齐)
            # DOCX 直接放 emoji, Word 在 Windows 系统会用 Segoe UI Emoji fallback 渲染
            _add_field('结构', _safe_text(getattr(item, 'structure_analysis', None)), icon='💡')
            _add_field('举一反三', _safe_text(getattr(item, 'similar_expressions', None)), icon='🔄')
            _add_field('场景', _safe_text(getattr(item, 'usage_scenario', None)), icon='🎯')
            _add_field('相似表达', _safe_text(getattr(item, 'alternative_phrasings', None)), icon='🔄')
            example_en = _safe_text(getattr(item, 'example_sentence', None))
            if example_en:
                _add_field('例句', example_en, icon='📝')
            ctx_cn = _safe_text(getattr(item, 'context_translation', None))
            if ctx_cn:
                pcn = cell.add_paragraph()
                _add_para_run(pcn, 'CN  ', bold=True, size=10, color_hex='94A3B8')
                _add_para_run(pcn, ctx_cn, size=10, color_hex='64748B')

            doc.add_paragraph('')  # 表达卡间距

    # === Phase E: A-Z 词汇索引 (DOCX-only, PDF 已 defer) ===
    # 在末尾页按首字母分组所有 words, 每个 word 是 hyperlink → 单词卡 section
    # 用户点击跳转后用 Word 的查找功能 (Ctrl+F) 定位具体单词
    word_list = [_safe_text(it.content_en) for it in groups['words'] if _safe_text(it.content_en)]
    # 去重保序
    seen = set()
    unique_words = []
    for w in word_list:
        if w.lower() not in seen:
            seen.add(w.lower())
            unique_words.append(w)
    _add_index_page(doc, unique_words)

    # === 页眉页脚 ===
    section = doc.sections[0]
    # 页眉
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # 右侧 tab stop 让视频名真右对齐 (页面宽 - margin = 178mm = 10070 twips)
    _add_para_run(header_p, 'Linyu 学习卡', size=9, color_hex='94A3B8')
    _add_para_run(header_p, '\t' + (material_title or '')[:40], size=9, color_hex='94A3B8')
    # 加右对齐 tab stop
    pPr_hdr = header_p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '10070')  # 178mm = 10070 twips (1mm ≈ 56.7 twips)
    tabs.append(tab)
    pPr_hdr.append(tabs)
    # 页眉底部分隔线 (对标 PDF canvas.line 0.4pt 灰)
    _set_paragraph_bottom_border(header_p, size_eighths=4, color='E2E8F0', space_pt=2)

    # 页脚: 使用域代码 PAGE
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_para_run(footer_p, 'From Linyu', size=9, color_hex='94A3B8')
    _add_para_run(footer_p, '\t— ', size=9, color_hex='94A3B8')
    # 插入页码域 (注意: 域代码 run 也得设字体大小, 否则会用默认)
    run = footer_p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string('94A3B8')
    run.font.name = 'Microsoft YaHei'
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
    rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
    rPr.append(rFonts)
    _add_para_run(footer_p, ' —', size=9, color_hex='94A3B8')
    # 页脚右侧 tab stop (让页码真右对齐)
    pPr_ftr = footer_p._p.get_or_add_pPr()
    tabs_ftr = OxmlElement('w:tabs')
    tab_ftr = OxmlElement('w:tab')
    tab_ftr.set(qn('w:val'), 'right')
    tab_ftr.set(qn('w:pos'), '10070')
    tabs_ftr.append(tab_ftr)
    pPr_ftr.append(tabs_ftr)

    # 序列化到 bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ==================== 文件名 (RFC 5987 编码, 中文不乱码) ====================

def build_filename(material_title: str, ext: str) -> str:
    """构造下载文件名: {视频标题}-学习卡片-{YYYY-MM-DD}.{ext}

    返回 URL-safe 的 filename* 形式 (RFC 5987)
    """
    today = datetime.date.today().strftime('%Y-%m-%d')
    title_clean = _sanitize_title(material_title) or '视频'
    raw = f'{title_clean}-学习卡片-{today}.{ext}'
    return raw
